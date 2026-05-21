from io import BytesIO
from typing import Any

from docx import Document


def generate_docx(payload: dict[str, Any]):
    buffer = BytesIO()
    doc = Document()

    title = payload.get("title") or "Quiz Download"
    quiz_type = payload.get("quiz_type") or "quiz"
    description = payload.get("description")

    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Type: {quiz_type}")
    if description:
        doc.add_paragraph(description)
    doc.add_heading("Questions", level=2)

    for item in payload.get("questions", []):
        question_number = item.get("number")
        question_prefix = f"Question {question_number}" if question_number else "Question"

        doc.add_paragraph(f"{question_prefix}: {item['question']}")
        for option in item.get("options") or []:
            doc.add_paragraph(option)

        doc.add_paragraph(f"Answer: {item['answer']}")
        doc.add_paragraph("")

    doc.save(buffer)
    buffer.seek(0)
    return buffer
